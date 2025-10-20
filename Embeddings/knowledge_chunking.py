import pandas as pd
import re
from docx import Document

doc = Document("Knowledge/Knowledge Base.docx")
paras = [p.text.strip() for p in doc.paragraphs]
tables = doc.tables

# Find the index of the first real chunk header (e.g., CNTR:, VSL:, etc.)
type_codes = ["CNTR", "VAS", "VSL", "EDI", "API", "EDIE", "ADI"]
header_pattern = re.compile(rf"^({'|'.join(type_codes)})(:|\s)", re.MULTILINE)
first_chunk_idx = next(i for i, p in enumerate(paras) if header_pattern.match(p))
paras = paras[first_chunk_idx:]

# Find all chunk start indices
chunk_indices = [i for i, p in enumerate(paras) if header_pattern.match(p)]

# --- 1. Create DataFrame for title, overview, rest_chunk ---
chunk_data = []
for idx, start in enumerate(chunk_indices):
    end = chunk_indices[idx+1] if idx+1 < len(chunk_indices) else len(paras)
    chunk_paras = paras[start:end]
    #title = chunk_paras[0].split(":", 1)[-1].strip() if ":" in chunk_paras[0] else chunk_paras[0].strip()
    title = chunk_paras[0].strip()

    # Extract overview
    overview = ""
    overview_start = None
    for i, line in enumerate(chunk_paras):
        if line.lower() == "overview":
            overview_start = i
            break
    if overview_start is not None:
        end_idx = None
        for i in range(overview_start+1, len(chunk_paras)):
            if chunk_paras[i].lower() in ["resolution", "preconditions", "decision logic", "verification"]:
                end_idx = i
                break
        if end_idx:
            overview = "\n".join(chunk_paras[overview_start+1:end_idx]).strip()
        else:
            overview = "\n".join(chunk_paras[overview_start+1:]).strip()

    rest_chunk = ""
    if overview_start is not None:
        rest_chunk = "\n".join(chunk_paras[overview_start+1:]).strip()

    chunk_data.append({
        "title": title,
        "overview": overview,
        "rest_chunk": rest_chunk
    })

df_chunks = pd.DataFrame(chunk_data)

# --- 2. Create DataFrame for modules from tables ---
module_data = []
for tbl in tables:
    found = False
    for row in tbl.rows:
        for cidx, cell in enumerate(row.cells):
            if cell.text.strip().lower() == "module":
                if cidx + 1 < len(row.cells):
                    module = row.cells[cidx + 1].text.strip()
                    module_data.append({"module": module})
                    found = True
                    break
        if found:
            break

df_modules = pd.DataFrame(module_data)

# --- 3. Combine the two DataFrames ---
# (Assume order matches: first module for first chunk, etc.)
df_combined = pd.concat([df_modules, df_chunks], axis=1)

print(df_combined)
df_combined.to_csv("knowledge_base_chunks_combined.csv", index=False)